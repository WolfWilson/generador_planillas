#!/usr/bin/env python
# coding: utf-8
"""
Generador de Planilla de Firmas (.docx)
- Lunes a viernes con horario de mañana
- Martes y jueves con extras por la tarde (configurable)
- Fines de semana marcados como SÁBADO/DOMINGO con guiones
- Permite marcar feriados o notas especiales por día
Requiere: python-docx  ->  pip install python-docx
"""

from __future__ import annotations
from docx import Document as _DocumentFactory
from docx.document import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import argparse
import calendar
from datetime import date, datetime
from typing import Iterable, Dict, Optional

def month_name_es(m: int, upper: bool = True) -> str:
    nombres = {
        1: "ENERO", 2: "FEBRERO", 3: "MARZO", 4: "ABRIL",
        5: "MAYO", 6: "JUNIO", 7: "JULIO", 8: "AGOSTO",
        9: "SEPTIEMBRE", 10: "OCTUBRE", 11: "NOVIEMBRE", 12: "DICIEMBRE"
    }
    s = nombres.get(m, "")
    return s if upper else s.capitalize()

def set_cell_shading(cell, fill="FFFFFF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_header(doc: DocxDocument, nombre: str, oficina: str, empleado: str, mes: int, anio: int):
    """Crea el encabezado con dos columnas (similar a la segunda captura)."""
    # Título centrado
    title = doc.add_paragraph("PLANILLAS PERSONAL DE HORARIOS")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    # Tabla 2x2 para alinear datos izquierda/derecha
    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = False
    try:
        tbl.columns[0].width = Cm(10)
        tbl.columns[1].width = Cm(7)
    except Exception:
        pass

    # Fila 0
    c00 = tbl.cell(0, 0).paragraphs[0]
    run = c00.add_run("APELLIDO Y NOMBRE: ")
    run.bold = True
    run.font.size = Pt(11)
    run = c00.add_run(nombre)
    run.font.size = Pt(11)

    c01 = tbl.cell(0, 1).paragraphs[0]
    run = c01.add_run("Oficina: ")
    run.bold = True
    run.font.size = Pt(11)
    run = c01.add_run(oficina)
    run.font.size = Pt(11)

    # Fila 1
    c10 = tbl.cell(1, 0).paragraphs[0]
    run = c10.add_run("Mes: ")
    run.bold = True
    run.font.size = Pt(11)
    run = c10.add_run(f"{month_name_es(mes, upper=True)} {anio}")
    run.font.size = Pt(11)

    c11 = tbl.cell(1, 1).paragraphs[0]
    run = c11.add_run("Empleado: ")
    run.bold = True
    run.font.size = Pt(11)
    run = c11.add_run(empleado)
    run.font.size = Pt(11)

    # Espacio después del encabezado
    doc.add_paragraph("")

def build_table(doc: DocxDocument, mes: int, anio: int,
                hora_maniana=("6:30", "13:00"),
                hora_tarde=("16:00", "19:00"),
                extras_dow: Iterable[int]=(1, 3),
                notas_por_dia: Optional[Dict[int, str]] = None,
                feriados: Optional[Iterable[date]] = None) -> None:
    notas_por_dia = notas_por_dia or {}
    feriados = set(feriados or [])

    rows = 2 + calendar.monthrange(anio, mes)[1]
    cols = 9
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    widths_cm = [1.3, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2]
    for i, w in enumerate(widths_cm):
        for r in range(rows):
            try:
                table.cell(r, i).width = Cm(w)
            except:
                pass

    hdr0 = table.rows[0]
    hdr0.cells[0].text = ""
    a = hdr0.cells[1]; b = hdr0.cells[4]
    a.merge(b); hdr0.cells[1].text = "MAÑANA"
    a2 = hdr0.cells[5]; b2 = hdr0.cells[8]
    a2.merge(b2); hdr0.cells[5].text = "TARDE"
    for j in [1, 5]:
        run = table.cell(0, j).paragraphs[0].runs[0]
        run.bold = True
        table.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    labels = ["DIAS", "ENTRADA", "FIRMA", "SALIDA", "FIRMA",
              "ENTRADA", "FIRMA", "SALIDA", "FIRMA"]
    for c, txt in enumerate(labels):
        cell = table.cell(1, c)
        cell.text = txt
        p = cell.paragraphs[0]
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "EFEFEF")

    ndays = calendar.monthrange(anio, mes)[1]
    for d in range(1, ndays + 1):
        r = 1 + d
        dow = calendar.weekday(anio, mes, d)  # 0=Mon ... 6=Sun
        is_weekend = dow >= 5
        current_date = date(anio, mes, d)

        table.cell(r, 0).text = str(d)
        table.cell(r, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        def fill_dashes():
            for c in range(1, 9):
                if c in (2, 4, 6, 8):
                    table.cell(r, c).text = ""
                else:
                    table.cell(r, c).text = "---"
                table.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        special_note = notas_por_dia.get(d)
        if is_weekend:
            label = "SÁBADO" if dow == 5 else "DOMINGO"
            table.cell(r, 1).text = label
            table.cell(r, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for c in range(2, 9):
                if c in (2, 4, 6, 8):
                    table.cell(r, c).text = ""
                else:
                    table.cell(r, c).text = "---"
                table.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for c in range(0,9):
                set_cell_shading(table.cell(r, c), "F8F8F8")
            continue

        if current_date in feriados:
            table.cell(r, 1).text = "FERIADO"
            table.cell(r, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_dashes()
            for c in range(0,9):
                set_cell_shading(table.cell(r, c), "FFF2CC")
            continue

        if special_note:
            table.cell(r, 1).text = special_note
            table.cell(r, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_dashes()
            continue

        table.cell(r, 1).text = hora_maniana[0]
        table.cell(r, 3).text = hora_maniana[1]
        for c in (1,3):
            table.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for c in (2,4,6,8):
            table.cell(r, c).text = ""

        if dow in extras_dow:
            table.cell(r, 5).text = hora_tarde[0]
            table.cell(r, 7).text = hora_tarde[1]
            for c in (5,7):
                table.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            table.cell(r, 5).text = "---"
            table.cell(r, 7).text = "---"
            for c in (5,7):
                table.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\\")

def generar_planilla_word(
    out_path: str,
    nombre: str,
    oficina: str,
    empleado: str,
    mes: int,
    anio: int,
    hora_maniana=("6:30", "13:00"),
    hora_tarde=("16:00", "19:00"),
    extras_dow: Iterable[int] = (1, 3),
    notas_por_dia: Optional[Dict[int, str]] = None,
    feriados: Optional[Iterable[date]] = None,
):
    # Crear nuevo documento (factoría) y luego usar la clase concreta para tipado.
    doc = _DocumentFactory()
    add_header(doc, nombre, oficina, empleado, mes, anio)
    doc.add_paragraph("")
    build_table(doc, mes, anio, hora_maniana, hora_tarde, extras_dow, notas_por_dia, feriados)
    doc.save(out_path)

def parse_date_list(text: str) -> set[date]:
    # "2025-07-09,2025-07-19"
    out = set()
    if not text:
        return out
    for tok in text.split(","):
        tok = tok.strip()
        try:
            y,m,d = map(int, tok.split("-"))
            out.add(date(y,m,d))
        except Exception:
            pass
    return out

def parse_notes(text: str) -> Dict[int, str]:
    # "16:LICENCIA DE INV.,17:LICENCIA DE INV."
    out: Dict[int, str] = {}
    if not text:
        return out
    for tok in text.split(","):
        tok = tok.strip()
        if ":" in tok:
            k, v = tok.split(":", 1)
            try:
                out[int(k)] = v.strip()
            except:
                continue
    return out

def main():
    ap = argparse.ArgumentParser(description="Generar planilla .docx mensual")
    ap.add_argument("--out", required=True, help="Ruta de salida .docx")
    ap.add_argument("--nombre", required=True, help="Apellido y nombre")
    ap.add_argument("--oficina", required=True, help="Oficina / Sector")
    ap.add_argument("--empleado", required=True, help="Legajo / DNI")
    ap.add_argument("--mes", type=int, required=True, help="Mes (1-12)")
    ap.add_argument("--anio", type=int, required=True, help="Año (e.g., 2025)")
    ap.add_argument("--hora_m", default="6:30,13:00", help="Horario mañana 'HH:MM,HH:MM'")
    ap.add_argument("--hora_t", default="16:00,19:00", help="Horario tarde 'HH:MM,HH:MM'")
    ap.add_argument("--extras_dow", default="1,3", help="Días con extras por la tarde (0=lun ... 6=dom). Por defecto: 1,3 (mar, jue)")
    ap.add_argument("--feriados", default="", help="Fechas feriado separadas por coma en formato YYYY-MM-DD")
    ap.add_argument("--notas", default="", help="Notas especiales por día: '16:LICENCIA DE INV.,17:CAPACITACIÓN'")

    args = ap.parse_args()

    hm = tuple(args.hora_m.split(","))
    ht = tuple(args.hora_t.split(","))
    extras = tuple(int(x) for x in args.extras_dow.split(",") if x.strip()!="")
    feriados = parse_date_list(args.feriados)
    notas = parse_notes(args.notas)

    generar_planilla_word(
        out_path=args.out,
        nombre=args.nombre,
        oficina=args.oficina,
        empleado=args.empleado,
        mes=args.mes,
        anio=args.anio,
        hora_maniana=hm, hora_tarde=ht,
        extras_dow=extras,
        notas_por_dia=notas,
        feriados=feriados,
    )

if __name__ == "__main__":
    main()
